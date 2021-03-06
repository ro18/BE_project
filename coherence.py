from flask import Blueprint, render_template
from difflib import SequenceMatcher
import nltk
import docx2txt
import pandas as pd
import docx
from docx import Document
import nltk as nl
import pandas as pd
import numpy as np
import operator
import glob
# coherence=Blueprint("coherence",__name__,static_folder="static",template_folder="templates")
# keywo = ["java", "python", "c"]
# @coherence.route("/coherence")
# def coherence_c():


def coherence_cv(keywo):
    print(keywo)
    keywords = keywo[0]
    for filename in glob.glob("./uploads/*.docx"):
        resume=filename
    my_text = docx2txt.process(resume)
    document = Document(resume)
    lst = pd.read_csv("./repository.csv")
    # a=my_text.encode("ISO-8859-1","ignore")
    # print(type(a))
    headings = []
    tables = document.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    headings.append(paragraph.text)
    headings = [str(k) for k in headings]
    sections = my_text.split('\n\n')
    sec_no = 0
    paras = ['']
    for sec in sections:
        if sec in headings:
            sec_no += 1
            paras.append('')
            continue
        paras[sec_no] += sec+"\n"
    tokens = []
    token_final = []
    for i, p in enumerate(paras):
        # print p
        tokens.append(nltk.tokenize.word_tokenize(p))
    #xy =nltk.tokenize.word_tokenize(paras)
    for token in tokens:
        token_final.append(token)
    paras_sent = []
    for para in paras:
        paras_sent.append([])
        paras_sent[-1].extend(nl.tokenize.sent_tokenize(para.replace('\n', '.')))
    paras_sent = []
    for para in paras:
        paras_sent.append([])
        paras_sent[-1].extend(para.split('\n'))
    repo = [str(list(k)[0]).lower() for k in list(np.array(lst))]
    #from  sets import Set

    def similar(my_text, b):
        return SequenceMatcher(None, my_text, b).ratio()

    repo = [str(list(k)[0]).lower() for k in list(np.array(lst))]

    personal = []
    for token in tokens:
        for retoken in token:
            for r in repo:
                x = retoken.lower()
                # print retoken
                if similar(x, r) > 0.9:
                    if r not in personal:
                        personal.append(r)
    my_text = my_text.lower()
    for i in range(0, len(repo)):
        st = str(repo[i]).lower()
        if my_text.find(st) >= 0:
            if repo[i] not in personal:
                personal.append(st)

    Technical_skill = personal
    print("The technical skills are")
    for i in range(0, len(Technical_skill)):
        print("{0}.{1}".format(i+1, Technical_skill[i]))

    my_text = my_text.replace(',', ' ')
    wordlist = my_text.split()
    wordfreq = [wordlist.count(w) for w in personal]
    # print("wordlist{}".format(wordlist))
    # print("wordfreq{}".format(wordfreq))
    # print("Pairs\n" + str(list(zip(personal, wordfreq))))

    d = dict(zip(personal, wordfreq))
    sorted_d = dict(
        sorted(d.items(), key=operator.itemgetter(1), reverse=True))
    # print(sorted_d)

    # keywords = [ "shell-scripting", "numpy","c"]
    print(keywords)
    print(type(keywords))
    print(keywords[0])
    print(keywords[1])
    print(keywords[2])
    a1 = [50, 45, 40, 35, 30, 25, 20, 15, 10, 7, 0]
    a2 = [30, 30, 28, 25, 22, 19, 16, 12, 8, 5, 0]
    a3 = [20, 20, 20, 18, 16, 14, 12, 10, 6, 3, 0]
    i = -1
    for w in sorted_d:
        i = i+1
        if (w == keywords[0]):
            print("one"+keywords[0])
            break
        if(i+1 == len(d)):
            i = 10
    j = -1
    for w in sorted_d:
        j = j+1
        if (w == keywords[1]):
            print("two"+keywords[1])
            break
        if(j+1 == len(d)):
            j = 10

    k = -1
    for w in sorted_d:
        k = k+1
        if (w == keywords[2]):
            break
        if(k+1 == len(d)):
            k = 10

    ans = a1[i]+a2[j]+a3[k]
    # print(ans
    print("after coherence cv")
    return ans,sorted_d


def coherence_ans(keywo):
    keywords = keywo[0]
    d = dict()
    my_text = open("./uploads/audio_text.txt", "r")
    with open("./repository.csv") as f:
        lst = f.read()
    lst = lst.replace('\n', ' ')
    # lst=lst.splitlines()
    lst = lst.lower()
    print(lst)
    wordz = lst.split(" ")
    wordz.pop()
    print(wordz)
    for line in my_text:
        line = line.strip()
        line = line.lower()
        words = line.split(" ")
        for word in words:
            if word in wordz:
                if word in d:
                    d[word] = d[word] + 1
                else:
                    d[word] = 1
    for key in list(d.keys()):
        print(key, ":", d[key])
    sorted_d = dict(
        sorted(d.items(), key=operator.itemgetter(1), reverse=True))
    print(sorted_d)
    # result = sorted_d.pop(" ", None)
    a1 = [50, 42, 35, 28, 21, 14, 7, 0]
    a2 = [30, 30, 25, 20, 15, 10, 5, 0]
    a3 = [20, 20, 20, 15, 11, 7, 3, 0]
    i = -1
    for w in sorted_d:
        i = i+1
        if (w == keywords[0]):
            break
        if(i+1 == len(d)):
            i = 7
    j = -1
    for w in sorted_d:
        j = j+1
        if (w == keywords[1]):
            break
        if(j+1 == len(d)):
            j = 7

    k = -1
    for w in sorted_d:
        k = k+1
        if (w == keywords[2]):
            break
        if(k+1 == len(d)):
            k = 7

    ans = a1[i]+a2[j]+a3[k]
    return ans,sorted_d

# ans1=coherence_cv(keywo)
# ans2=coherence_ans(keywo)

# print(ans1)
# print(ans2)
